from fastapi import HTTPException
from typing import List, Dict
from app.core.database import get_percieve_db_connection
from app.core.azure_client import client

from app.services.ip_validity_analysis.novelty import compare_novelty
from app.services.ip_validity_analysis.non_obviousness import compare_non_obviousness
from app.services.ip_validity_analysis.utility import compare_utility
from app.services.ip_validity_analysis.enablement import compare_enablement
from app.services.ip_validity_analysis.written_description import (
    compare_written_description,
)
from app.services.ip_validity_analysis.definiteness import compare_definiteness
from app.services.ip_validity_analysis.industrial_application import (
    compare_industrial_application,
)
from app.services.ip_validity_analysis.clarity_and_sufficiency import (
    compare_clarity_and_sufficiency,
)
from app.services.ip_validity_analysis.scope_and_definition import (
    compare_scope_and_definition,
)
from app.services.ip_validity_analysis.economic_significance import (
    compare_economic_significance,
)
from app.models.ip_validity_analysis import (
    PatentAnalysis,
    PatentResult,
    PatentList,
    SearchQuery,
    ReportParams,
)
from psycopg2 import sql
from sklearn.metrics.pairwise import cosine_similarity
from app.utils.prior_art_search_helpers import (
    vectorize_description,
    vectorize_description_with_retries,
)
import time
import numpy as np
from app.services.add_attachment_answer import get_report_id
from app.core.database import get_db_connection
import json

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import os
import base64
from io import BytesIO


async def search_documents(keywords: List[str]) -> List[PatentResult]:
    try:
        conn = get_percieve_db_connection()
        query_keywords = " | ".join(keywords)
        ts_query = sql.SQL("plainto_tsquery('english', %s)")

        query = sql.SQL(
            """
            SELECT id, title, COALESCE(abstract, ''), claim_text
            FROM target.patents
            WHERE ((title_abstract_tsvector @@ {ts_query} AND title IS NOT NULL AND title != '' AND abstract IS NOT NULL AND abstract != '' AND claim_text IS NOT NULL AND claim_text != '')
            OR (claims_tsvector @@ {ts_query} AND title IS NOT NULL AND title != '' AND abstract IS NOT NULL AND abstract != '' AND claim_text IS NOT NULL AND claim_text != ''));
            """
        ).format(ts_query=ts_query)

        with conn.cursor() as cur:
            cur.execute(query, [query_keywords, query_keywords])
            results = cur.fetchall()

            if not results:
                print("No patents found.")
                return []

            patents = []
            for row in results:
                try:
                    id_, title, abstract, content = row[:4]
                    patents.append(
                        PatentResult(
                            id=id_,
                            title=title,
                            abstract=abstract,  # Use COALESCE to handle NULL or empty abstract values
                            content=content,
                        )
                    )
                except Exception as e:
                    print(f"Error processing row: {row}")
                    print(f"Error details: {e}")
                    continue

            return patents
    except Exception as e:
        print(f"Error executing search query: {e}")
        raise HTTPException(status_code=500, detail="Error executing search query")
    finally:
        conn.close()


def encode_texts(texts: List[str]) -> List[List[float]]:
    embeddings = vectorize_texts(texts)
    print("Encoding summary process completed.")
    return embeddings


def vectorize_texts(texts: List[str]):
    embeddings = []
    for text in texts:
        try:
            print(f"Vectorizing text: {text}")
            embedding = vectorize_description_with_retries(text)
            embeddings.append(embedding)
            print(f"Successfully vectorized text: {text}")
            # Add a 10 second delay
            time.sleep(10)
        except Exception as e:
            print(f"Failed to vectorize text: {text}. Error: {e}")

    return embeddings


async def search_patents(patents: List[PatentResult], description: str) -> PatentList:
    patent_texts = [f"{patent.abstract} " for patent in patents]

    # Ensure that you await the results of the asynchronous tasks
    description_embedding_task = encode_texts([description])
    patent_embeddings_task = vectorize_texts(patent_texts)

    description_embedding = description_embedding_task
    patent_embeddings = patent_embeddings_task

    # description_embedding = np.array(description_embedding[0])
    # patent_embeddings = np.array(patent_embeddings)

    # Convert lists to NumPy arrays
    description_embedding = np.array(description_embedding).reshape(1, -1)
    patent_embeddings = np.array(patent_embeddings)

    # Ensure patent_embeddings are reshaped correctly
    if len(patent_embeddings.shape) == 1:
        patent_embeddings = patent_embeddings.reshape(1, -1)

    # Compute cosine similarity
    similarities = cosine_similarity(description_embedding, patent_embeddings)

    # Print similarities for each patent
    similarity_dict: Dict[str, float] = {}
    for i, patent in enumerate(patents):
        similarity = similarities[0][i]
        similarity_dict[patent.id] = similarity
        print(f"Similarity with patent {patent.id}: {similarity}")

    # Sort the similarities and get the top 10 most similar patent IDs
    sorted_similarities = sorted(
        similarity_dict.items(), key=lambda item: item[1], reverse=True
    )
    top_10_patent_ids = [patent_id for patent_id, _ in sorted_similarities[:5]]

    return top_10_patent_ids


async def create_assessment(patent_ids: List[str], answer_list: str, criterion: str):
    patent_abstracts = []

    print("Patent IDs: ", patent_ids)
    for patent_id in patent_ids:
        patent = await get_patent_by_id(patent_id)
        if patent:
            patent_abstracts.append(patent)

    if not patent_abstracts:
        raise HTTPException(
            status_code=404, detail="No patent abstracts found for the provided IDs."
        )

    if criterion == "Novelty (35 U.S.C. § 102)":
        assessment_point = await compare_novelty(patent_abstracts, answer_list)
    elif criterion == "Non-Obviousness (35 U.S.C. § 103)":
        assessment_point = await compare_non_obviousness(patent_abstracts, answer_list)
    elif criterion == "Utility (35 U.S.C. § 101)":
        assessment_point = await compare_utility(patent_abstracts, answer_list)
    elif criterion == "Enablement (35 U.S.C. § 112(a))":
        assessment_point = await compare_enablement(patent_abstracts, answer_list)
    elif criterion == "Written Description (35 U.S.C. § 112(a))":
        assessment_point = await compare_written_description(
            patent_abstracts, answer_list
        )
    elif criterion == "Definiteness (35 U.S.C. § 112(b))":
        assessment_point = await compare_definiteness(patent_abstracts, answer_list)
    elif criterion == "Industrial Application":
        assessment_point = await compare_industrial_application(
            patent_abstracts, answer_list
        )
    elif criterion == "Clarity & Sufficiency":
        assessment_point = await compare_clarity_and_sufficiency(
            patent_abstracts, answer_list
        )
    elif criterion == "Scope & Definition":
        assessment_point = await compare_scope_and_definition(
            patent_abstracts, answer_list
        )
    elif criterion == "Economic Significance":
        assessment_point = await compare_economic_significance(
            patent_abstracts, answer_list
        )

    return assessment_point


async def get_patent_by_id(patent_id: str):
    query = """
    SELECT abstract
    FROM target.patents
    WHERE id = %s;
    """
    value = patent_id
    try:
        conn = get_percieve_db_connection()

        with conn.cursor() as cur:
            cur.execute(query, [value])
            result = cur.fetchone()
            return result[0] if result else None
    except Exception as e:
        print(f"Error fetching patent {patent_id}: {e}")
        return None
    finally:
        conn.close()


async def get_answers(requirement_gathering_id, user_case_id):
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Check if requirement_gathering_id exists in user_chats
        query_user_chats = """
        SELECT COUNT(*)
        FROM user_chats
        WHERE requirement_gathering_id = %s;
        """

        print("1")
        cur.execute(query_user_chats, (requirement_gathering_id,))
        user_chats_count = cur.fetchone()[0]

        # Check if requirement_gathering_id exists in attachment_chats
        query_attachment_chats = """
        SELECT COUNT(*)
        FROM attachment_chats
        WHERE requirement_gathering_id = %s;
        """
        cur.execute(query_attachment_chats, (requirement_gathering_id,))
        attachment_chats_count = cur.fetchone()[0]

        if user_chats_count > 0:

            question_ids = []
            # Determine question_ids based on user_case_id
            if user_case_id == "1" or user_case_id == "2":
                question_ids = [
                    "1",
                    "2",
                    "6",
                    "7",
                    "8",
                    "9",
                    "10",
                    "11",
                    "12",
                    "13",
                ]
            elif user_case_id == "3":
                question_ids = [
                    "14",
                    "15",
                    "16",
                    "17",
                    "18",
                    "19",
                    "20",
                    "21",
                    "22",
                    "23",
                    "24",
                    "25",
                    "26",
                ]
            elif user_case_id == "4":
                question_ids = ["27", "28", "29", "30", "31", "32", "33", "34", "35"]
            elif user_case_id == "5":
                question_ids = ["36", "37", "38", "39", "40", "41", "42"]
            elif user_case_id == "6":
                question_ids = [
                    "1",
                    "2",
                    "3",
                    "4",
                    "43",
                    "44",
                    "45",
                    "46",
                    "47",
                    "48",
                    "49",
                    "50",
                    "51",
                    "52",
                    "53",
                    "54",
                    "55",
                ]
            elif user_case_id == "7":
                question_ids = [
                    "1",
                    "2",
                    "3",
                    "4",
                    "56",
                    "57",
                    "58",
                    "59",
                    "60",
                    "61",
                ]
            elif user_case_id == "8":
                question_ids = [
                    "1",
                    "2",
                    "3",
                    "4",
                    "62",
                    "63",
                    "64",
                    "65",
                    "66",
                    "67",
                    "68",
                    "69",
                    "70",
                    "71",
                    "72",
                ]
            elif user_case_id == "9":
                question_ids = [
                    "1",
                    "2",
                    "3",
                    "4",
                    "73",
                    "74",
                    "75",
                    "76",
                    "77",
                    "78",
                    "79",
                    "80",
                ]
            elif user_case_id == "10":
                question_ids = [
                    "1",
                    "2",
                    "3",
                    "4",
                    "81",
                    "82",
                    "83",
                    "84",
                    "85",
                    "86",
                    "87",
                    "88",
                    "89",
                ]
            else:
                question_ids = ["0", "1", "2"]
            query = """
            SELECT answer
            FROM user_chats
            WHERE requirement_gathering_id = %s AND question_id IN %s;
            """
            values = (requirement_gathering_id, tuple(question_ids))

        elif attachment_chats_count > 0:

            report_id = await get_report_id(requirement_gathering_id, user_case_id)
            query = """
            SELECT answer
            FROM attachment_chats
            WHERE requirement_gathering_id = %s AND report_id = %s;
            """
            values = (
                requirement_gathering_id,
                report_id,
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid type_id")
        cur.execute(query, values)
        result = cur.fetchall()

        answers = [row[0] for row in result]
        print(f"Number of answers found: {len(answers)}")

        # Combine answers into a single paragraph
        paragraph = " ".join(answers)
        return paragraph
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn:  # Ensure conn is closed only if it was successfully opened
            conn.close()


async def get_keywords(answers):
    system_prompt = f"Extract exactly 2 general product-related keywords from the following text. Ensure these are broad terms like 'satellite' or 'motor' and not specific names, companies, or places. Separate them with a comma: {answers}"
    try:
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": system_prompt},
        ]

        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            temperature=0.7,
            max_tokens=800,
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
            stop=None,
        )
        content = completion.choices[0].message.content
        print(content)

        # Extract and clean keywords
        if content:
            keywords = [keyword.strip() for keyword in content.split(",") if keyword]
            return keywords
        else:
            return {"status": "No content received"}
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error generating response: {str(e)}"
        )


async def add_report(report_str, requirement_gathering_id, user_case_id):
    query = """
    INSERT INTO reports (requirement_gathering_id, user_case_id, report)
    VALUES (%s, %s, %s)
    """
    values = (requirement_gathering_id, user_case_id, report_str)

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(query, values)

        conn.commit()  # Commit only after both operations succeed
        cur.close()
    except Exception as e:
        conn.rollback()  # Rollback if any exception occurs
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


async def get_summary(answers):
    system_prompt = f"Summarize the following text: {answers}"
    try:
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": system_prompt},
        ]

        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            temperature=0.7,
            max_tokens=800,
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
            stop=None,
        )

        content = completion.choices[0].message.content

        return content
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error generating response: {str(e)}"
        )


async def create_report_background(report_params: ReportParams):
    try:
        answers = await get_answers(
            report_params.requirement_gathering_id, report_params.user_case_id
        )
        summary = await get_summary(answers)
        keywords = await get_keywords(answers)
        response_data = await search_documents(keywords)
        response_data = response_data[:3]
        response_data2 = await search_patents(response_data, summary)

        patentability_criteria = [
            "Novelty (35 U.S.C. § 102)",
            "Non-Obviousness (35 U.S.C. § 103)",
            "Utility (35 U.S.C. § 101)",
            "Enablement (35 U.S.C. § 112(a))",
            "Written Description (35 U.S.C. § 112(a))",
            "Definiteness (35 U.S.C. § 112(b))",
            "Industrial Application",
            "Clarity & Sufficiency",
            "Scope & Definition",
            "Economic Significance",
        ]

        report = {}

        for i in range(len(patentability_criteria)):
            assessment = await create_assessment(
                response_data2,
                answers,
                patentability_criteria[i],
            )
            report[patentability_criteria[i]] = assessment

        # Convert report dictionary to JSON string
        # report_str = str(report)
        report_str = dict_to_formatted_string(report)

        await create_word_document(
            report_str,
            report_params.requirement_gathering_id,
            report_params.user_case_id,
        )
        await create_pdf_document(
            report_str,
            report_params.requirement_gathering_id,
            report_params.user_case_id,
        )

        await add_report(
            report_str,
            report_params.requirement_gathering_id,
            report_params.user_case_id,
        )
        print("Report generated successfully.")
        return report
    except Exception as e:
        query_status = """
        UPDATE report_file_status
        SET status = 'failed',description = %s
        WHERE requirement_gathering_id = %s AND use_case_id = %s;
        """
        values_status = (
            e,
            report_params.requirement_gathering_id,
            report_params.user_case_id,
        )
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute(query_status, values_status)
            conn.commit()
            cur.close()
        except Exception as e:
            conn.rollback()
            raise HTTPException(status_code=500, detail=str(e))

        raise HTTPException(status_code=500, detail=str(e))


def dict_to_formatted_string(report: Dict) -> str:
    report_str = ""
    for key, value in report.items():
        report_str += f"{key}:\n"
        if isinstance(value, dict):
            for sub_key, sub_value in value.items():
                report_str += f"  {sub_key}: {sub_value}\n"
        else:
            report_str += f"  {value}\n"
        report_str += "\n"
    return report_str


async def create_pdf_document(
    content: str, requirement_gathering_id: str, user_case_id: str
):
    buffer = BytesIO()

    # Create the PDF document
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica", 12)
    text = c.beginText(40, height - 40)

    paragraphs = content.split("\n")
    for para in paragraphs:
        text.textLines(para)
        text.textLine("")

    c.drawText(text)
    c.save()

    # Move to the beginning of the BytesIO buffer
    buffer.seek(0)

    # Convert the PDF to base64
    pdf_bytes = buffer.read()
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")

    file_type = "pdf"
    file_name = "IP Validity Analysis Report"

    # Insert the base64 PDF into the database
    await insert_file_to_db(
        pdf_base64,
        requirement_gathering_id,
        user_case_id,
        file_type,
        file_name,
        content,
    )


async def create_word_document(
    content: str, requirement_gathering_id: str, user_case_id: str
):
    buffer = BytesIO()

    doc = Document()
    doc.add_heading("Patent Novelty Assessment", 0)

    paragraphs = content.split("\n")
    for para in paragraphs:
        if para.strip() != "":
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.style.font.size = Pt(12)

    doc.save(buffer)

    # Move to the beginning of the BytesIO buffer
    buffer.seek(0)

    # Convert the Word document to base64
    word_bytes = buffer.read()
    word_base64 = base64.b64encode(word_bytes).decode("utf-8")

    file_type = "docx"
    file_name = "IP Validity Analysis Report"

    # Insert the base64 Word document into the database
    await insert_file_to_db(
        word_base64,
        requirement_gathering_id,
        user_case_id,
        file_type,
        file_name,
        content,
    )


async def insert_file_to_db(
    pdf_base64: str,
    requirement_gathering_id: str,
    user_case_id: str,
    file_type: str,
    file_name: str,
    content: str,
):
    query_status = """
     SELECT index FROM report_file_status WHERE requirement_gathering_id = %s AND use_case_id = %s;
    """
    values_status = (requirement_gathering_id, user_case_id)

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(query_status, values_status)
        status_index = cur.fetchone()
        cur.close()

        if not status_index:
            raise HTTPException(status_code=404, detail="Report file status not found.")

        query = """
        INSERT INTO report_file (requirement_gathering_id, use_case_id, file, file_type, file_name, report_file_status_id)
        VALUES (%s, %s, %s, %s, %s, %s)
        ON CONFLICT (requirement_gathering_id, use_case_id, file_type)
        DO UPDATE SET
            file = EXCLUDED.file,
            file_type = EXCLUDED.file_type,
            file_name = EXCLUDED.file_name,
            report_file_status_id = EXCLUDED.report_file_status_id
        """
        values = (
            requirement_gathering_id,
            user_case_id,
            pdf_base64,
            file_type,
            file_name,
            status_index[0],
        )

        query_status = """
        UPDATE report_file_status
        SET status = 'completed',
        description = %s
        WHERE requirement_gathering_id = %s AND use_case_id = %s;
        """
        values_status = (content, requirement_gathering_id, user_case_id)

        cur = conn.cursor()
        cur.execute(query, values)
        cur.execute(query_status, values_status)
        conn.commit()
        cur.close()
    except Exception as e:
        conn.rollback()
        raise HTTPException(
            status_code=500, detail=f"Database operation failed: {str(e)}"
        )
    finally:
        conn.close()


async def get_report_file(
    requirement_gathering_id: int, user_case_id: str, file_type: str
):
    query = """
    SELECT file, file_type,file_name
    FROM report_file
    WHERE requirement_gathering_id = %s AND use_case_id = %s AND file_type = %s;
    """
    values = (requirement_gathering_id, user_case_id, file_type)
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(query, values)
        result = cur.fetchone()
        if result:
            file_base64 = base64.b64encode(result[0]).decode("utf-8")
            report = {
                "file": file_base64,
                "file_type": result[1],
                "file_name": result[2],
            }
            return report
        else:
            raise HTTPException(status_code=404, detail="Report not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
        print(f"Error fetching report: {e}")
    finally:
        conn.close()
