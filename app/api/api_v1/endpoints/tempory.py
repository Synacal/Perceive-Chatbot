from urllib import request
from fastapi import APIRouter, HTTPException
from app.services.check_user_answers import check_user_answers
from psycopg2.extras import execute_values
import json
from app.core.azure_client import client

import pdfkit
import markdown
import os
import random
from bs4 import BeautifulSoup
from docx import Document
import matplotlib.pyplot as plt
import re
from dotenv import load_dotenv

from xhtml2pdf import pisa

load_dotenv()






number_set = set()

config=pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')
pdf_options = {
    'page-size': 'A4',
    'encoding': 'UTF-8',
}
system_prompt= """
Answer the following questions to the best of your ability. If you do not know the answer, Answer should be more than 1000 words. please type 'I do not know
Please generate a detailed document using the following structure and formatting requirements:

1. **Title**: The main title of the document should be formatted as a level 1 heading and use the font "Courier New, monospace".

2. **Subtopics**: Each main section should be formatted as a level 2 heading and use the font "Comic Sans MS, cursive". Sub-sections under each main section should be formatted as level 3 headings and use the font "Georgia, serif".

3. **Bold Text**: Emphasize important keywords or phrases by making them bold.

4. **Italic Text**: Use italics for additional emphasis, quotes, or specific terms that need to stand out.

5. **Fonts**: Use different fonts for different elements:
   - Main title: "Courier New, monospace"
   - Subtopics: "Comic Sans MS, cursive"
   - Sub-subtopics: "Georgia, serif"

Here is an example structure to follow:

# Main Title

<style>
h1 { font-family: 'Courier New', monospace; }
h2 { font-family: 'Comic Sans MS', cursive; }
h3 { font-family: 'Georgia', serif; }
</style>

## Subtopic 1

This is an introductory paragraph for Subtopic 1. It contains an **important keyword** and some *italicized text*.

### Sub-subtopic 1.1

This section goes into more detail under Subtopic 1. It also includes some **bold** and *italic** text for emphasis.

## Subtopic 2

This is an introductory paragraph for Subtopic 2. Similar to the previous sections, it can have **bold** keywords and *italicized** phrases.

### Sub-subtopic 2.1

Additional details for Subtopic 2. Use **bold** text for highlighting and *italic** text for emphasis.

"""

system_prompt_for_graph="""When a user asks to create a graph, provide the data strictly in the format 'x values and y values , chart title , x axis title and y axis title' without additional information.this is very importatnt. dont say any other things. give the answer in only this structure. refer the below examples to generate your answer. For example:

For the last five years of Apple product sales draw a bar chart: '2018: 868586, 2019: 6576845, 2020: 457843758, 2021: 8687475786, 2022: 876854 @Sales of apple product last five years @Year @products'.
For electric vehicle sales over the last 5 years draw a graph: '2017: 198350, 2018: 361307, 2019: 531407, 2020: 752684, 2021: 1051749 @Electric vehicle sales for last five years @Year @sales'.
For the population growth in major cities over the last decade draw a column chart: '2013: 1200000, 2014: 1350000, 2015: 1500000, 2016: 1650000, 2017: 1800000 @Population growth in major cities over the last decade @Year @Population'.
For the quarterly revenue of a company in the current fiscal year: 'Q1: 5000000, Q2: 5500000, Q3: 6000000, Q4: 6500000 @Quarterly revenue of a company in the current fiscal year @Year @Quarterly revenue'.
Please ensure that only these specific formats are provided without any additional explanations or details.dont generate any other words. only data set"""
system_prompt_mod="""Generate a detailed document using the following structure and formatting requirements:

1. **Title**: The main title of the document should be formatted as a level 1 heading and use the font "Courier New, monospace".
2. **Subtopics**: Each main section should be formatted as a level 2 heading and use the font "Comic Sans MS, cursive". Sub-sections under each main section should be formatted as level 3 headings and use the font "Georgia, serif".
3. **Bold Text**: Emphasize important keywords or phrases by making them bold.
4. **Italic Text**: Use italics for additional emphasis, quotes, or specific terms that need to stand out.
5. **Lists**: Use both ordered and unordered lists where appropriate to organize information.
6. **Links**: Include hyperlinks to relevant resources.
7. **Images**: Add images to illustrate key points or examples.
8. **Code Blocks**: Use code blocks for any programming-related examples or explanations.
9. **Tables**: Create tables to present structured data.
10. **Blockquotes**: Use blockquotes for notable quotes or important points.
11. **Horizontal Rules**: Use horizontal rules to separate sections or topics.
12. **Task Lists**: Include task lists if the content involves steps or checklists.

If you do not know the answer to a question or cannot provide detailed information, clearly state 'I do not know.' Ensure that the document is detailed and exceeds 1000 words.

Here is an example structure to follow:


<style>
h1 { font-family: 'Courier New', monospace; }
h2 { font-family: 'Comic Sans MS', cursive; }
h3 { font-family: 'Georgia', serif; }
</style>
# Main Title
## Subtopic 1

    This is an introductory paragraph for Subtopic 1. It contains an **important keyword** and some *italicized text*.

        - **Bold List Item 1**
        - *Italicized List Item 2*

### Sub-subtopic 1.1

            This section goes into more detail under Subtopic 1. It also includes some **bold** and *italic* text for emphasis.

            \`\`\`python
            # Sample Code Block
            def example_function():
                print("Hello, World!")
            \`\`\`

            [Link to Resource](https://www.example.com)

            ![Alt Text](image-url.jpg)

            > This is a blockquote highlighting a significant point.

## Subtopic 2

    This is an introductory paragraph for Subtopic 2. Similar to the previous sections, it can have **bold** keywords and *italicized* phrases.

        ### Sub-subtopic 2.1

        Additional details for Subtopic 2. Use **bold** text for highlighting and *italic* text for emphasis.

        | Column 1 | Column 2 |
        |----------|----------|
        | Row 1    | Data 1   |
        | Row 2    | Data 2   |

        ---

        - [ ] Task 1
        - [x] Task 2



"""
def generate_unique_number():
    while True:
        random_integer = random.randint(10000000, 20000000)
        if random_integer not in number_set:
            number_set.add(random_integer)
            return random_integer

def create_pdf(content,save_path):
    markdown_content = markdown.markdown(content)
    html_content = f"<html><body>{markdown_content}</body></html>"
       
    with open('temp.html', 'w', encoding='utf-8') as f:
        f.write(html_content)    

    pdfkit.from_file('temp.html', save_path, configuration=config,options=pdf_options)
    os.remove('temp.html')

def create_doc(content, save_path):
    doc = Document()
    markdown_content = markdown.markdown(content)
    html_content = f"<html><body>{markdown_content}</body></html>"

    soup = BeautifulSoup(html_content, 'html.parser')

    def add_elements_from_html(element):
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'b':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
        elif element.name == 'i':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.italic = True
        elif element.name == 'u':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.underline = True

    for tag in soup.body:
        if isinstance(tag, str):  
            continue
        add_elements_from_html(tag)
    doc.save(save_path)
def convert_to_dic(content):
    
    try:
       
        data_str, title, x_axis_title, y_axis_title = content.rsplit('@', 3)
    except ValueError:
        raise ValueError("Content format is incorrect. Expected format is 'data @title @x_axis_title @y_axis_title'.")

    split_data = data_str.split(',')
    data_dic = {}
    
    for word in split_data:
        try:
            X, Y = word.split(':')
            X = X.strip().strip("'")
            Y = Y.strip().strip("'")

            if 'million' in Y:
                Y = int(re.sub(r'[^0-9]', '', Y)) * 1_000_000
            elif 'billion' in Y:
                Y = int(re.sub(r'[^0-9]', '', Y)) * 1_000_000_000
            else:
                Y = int(re.sub(r'[^0-9]', '', Y))

            data_dic[int(X)] = Y

        except ValueError as e:
            print(f"Error parsing '{word}': {e}")
            raise ValueError(f"Error parsing '{word}': {e}")

    title = title.strip().strip("'")
    x_axis_title = x_axis_title.strip().strip("'")
    y_axis_title = y_axis_title.strip().strip("'")

    return data_dic, title, x_axis_title, y_axis_title

def generate_image_bar_chart(data_dic, title, x_axis_title, y_axis_title):
    plt.figure(figsize=(8, 5))  
    plt.bar(data_dic.keys(), data_dic.values(), color='skyblue')
    plt.xlabel(x_axis_title.strip("'"))
    plt.ylabel(y_axis_title.strip("'"))
    plt.title(title)
    plt.grid(True)
    plt.xticks(list(data_dic.keys())) 
    plt.tight_layout()

    # Save the plot as PNG image
    plt.savefig('bar_chart.png', dpi=300) 
    plt.close()

def generate_image_pie_chart(data_dic, title):
    plt.figure(figsize=(8, 5))
    plt.pie(data_dic.values(), labels=data_dic.keys(), autopct='%1.1f%%', startangle=140)
    plt.title(title.strip("'"))
    plt.tight_layout()

    # Save the plot as PNG image
    plt.savefig('pie_chart.png', dpi=300)
    plt.close()

def generate_image_line_chart(data_dic, title, x_axis_title, y_axis_title):
    plt.figure(figsize=(8, 5))  
    plt.plot(data_dic.keys(), data_dic.values(), marker='o')
    plt.xlabel(x_axis_title.strip("'"))
    plt.ylabel(y_axis_title.strip("'"))
    plt.title(title)
    plt.grid(True)
    plt.xticks(list(data_dic.keys())) 
    plt.tight_layout()

    # Save the plot as PNG image
    plt.savefig('line_chart.png', dpi=300) 
    plt.close()

def convert_html_to_pdf(html_string, pdf_path):
    with open(pdf_path, "wb") as pdf_file:
        pisa_status = pisa.CreatePDF(html_string, dest=pdf_file)
        
    return not pisa_status.err
def convert_to_html(content):
        markdown_content = markdown.markdown(content)
        html_content = f"<html><body>{markdown_content}</body></html>"
        return html_content



router = APIRouter()


@router.get("/tempory/")
async def tempory():
    
    user_prompt = "what is the difference between merge sort and quick sort?"
    message_text = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    try:
        completion = client.chat.completions.create(
            model="gpt-35-turbo",
            messages=message_text,
            temperature=0.7,
            max_tokens=800,
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
            stop=None,
        )

        content = f"{completion.choices[0].message.content}"
        
        unique_integer =generate_unique_number()
        
        save_folder = r'E:\intern\Synacal\New folder\Perceive-Chatbot\app\api\api_v1\endpoints\save_files'
        os.makedirs(save_folder, exist_ok=True)
        save_pdf_path = os.path.join(save_folder, f'file{unique_integer}.pdf')

        #create pdf
        html_content = convert_to_html(content)
        if convert_html_to_pdf(html_content, save_pdf_path):
            print(f"PDF generated and saved at {save_pdf_path}")
        else:
            print("PDF generation failed")

        # save_doc_path = os.path.join(save_folder, f'file{unique_integer}.docx')
       
        # create_pdf(content,save_pdf_path)
        # create_doc(content,save_doc_path)


        #create pdf using wkhtmltopdf
        file_name='test.pdf'
        create_pdf(content,file_name)

        #bar chart
        # try:
        #     data_dic,title,x_axis_title,y_axis_title = convert_to_dic(content)
        #     generate_image_bar_chart(data_dic,title,x_axis_title,y_axis_title)
        #     print("Image generated successfully.")
        # except Exception as e:
        #     print(f"An error occurred: {e}")

        # #pie chart
        # try:
        #     data_dic,title,x_axis_title,y_axis_title = convert_to_dic(content)
        #     generate_image_pie_chart(data_dic,title)
        #     print("Image generated successfully.")
        # except Exception as e:
        #     print(f"An error occurred: {e}")

        # #line chart
        # try:
        #     data_dic,title,x_axis_title,y_axis_title = convert_to_dic(content)
        #     generate_image_line_chart(data_dic,title,x_axis_title,y_axis_title)
        #     print("Image generated successfully.")
        # except Exception as e:
        #     print(f"An error occurred: {e}")


        
        


        
        return {"response": content}
    except Exception as e:
        return {"error": str(e)}
