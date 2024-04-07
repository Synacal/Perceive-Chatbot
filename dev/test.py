from docx import Document

def create_questionnaire(questions, file_path):
    # Create a new Document
    doc = Document()
    doc.add_heading('Questionnaire', 0)

    # Add questions to the document with numbering
    for i, question in enumerate(questions, start=1):
        doc.add_paragraph(f'{i}. {question}')

    # Save the document
    doc.save(file_path)

# List of common questions
questions = [
                #Common Qs
                'What is the full name of the company, and what is its core mission?',
                'Please provide a concise description of the technology or product your company has developed.',
                'Describe the technical aspects and unique features of the key product or technology developed by the company. How does your product or technology introduce innovation or novelty within its field?',
                "Could you explain your company's business model and how it generates revenue? What are the different revenue streams for your company, including primary and potential ancillary streams?",
                'What is your strategy for protecting the intellectual property associated with your product or technology? Are there specific patents or prior art that you have encountered during your research? What similarities or differences did you find?',
                
                #Synthetic data for IP validity analysis 
                # 'What is the full name of the company developing the key product or technology developed by the company?',
                # 'Please provide a concise description of the key product or technology developed by the company technology.',
                'Can you tell me more about the specific patents or prior art you may have encountered during your research? What similarities or differences did you find?',
                'How does the key product or technology developed by the company meet the criteria of novelty in its field?',
                'Can you explain why the features of the key product or technology developed by the company are considered non-obvious to someone skilled in the field?',
                'How is the key product or technology developed by the company applicable to industrial needs in its domain?',
                'What is your strategy for patent filing, including geographies and patent offices?',
                'How have you ensured enablement in the patent application for the key product or technology developed by the company?',
                'How have you ensured the definiteness of claims in your patent application for the key product or technology developed by the company?',
                'Can you provide the exact claims that will be present in the patent application for your key product or technology developed?',

                #IP licensing strategy process document
                'What specific technologies or innovations within NeuraWear are you looking to license, and what makes these aspects unique and valuable for potential licensees?',
                'Who are your ideal licensees for NeuraWear\'s technology, and in which industries or sectors do they primarily operate?',
                'What business goals are you aiming to achieve through IP licensing?',
                'What is your preferred licensing model for NeuraWear, and how does this preference align with your strategic objectives?',
                'Are there specific geographic regions you are targeting for licensing NeuraWear\'s technology?',
                'What are your financial expectations from licensing agreements?',
                'How prepared are you to negotiate and manage complex licensing agreements?',
                'What key terms and conditions are you prioritizing in your licensing agreements?',
                'Are you open to exploring strategic partnerships or cross-licensing opportunities?',
                'What metrics and KPIs will you use to evaluate the success of your licensing strategy?',
                'Do you have any performance requirements or specific expectations from licensees to ensure they contribute effectively to the licensed technology\'s success?',
                'How do you plan to handle sublicensing rights, audit rights, and quality control provisions to safeguard the integrity and value of your licensed IP?',
                'Are there any particular fields of use you are considering for NeuraWear\'s licensing agreements, and how do these choices reflect market demands and opportunities?',
                
                #IP Valuation questions list
                'What is the pricing strategy for your product or service?',
                'How do you calculate the gross margin for your offerings?',
                'What are the total development costs incurred for your product or service?',
                'What future costs do you anticipate for full development and market launch?',
                'What discount rate do you apply to future cash flows and why?',
                'What is the projected annual revenue growth rate, and how did you arrive at this figure?',
                'What are the anticipated operating expenses, and how are they allocated?',
                'How do you project sales revenue for your products or services over the next 5 years?',
                'What market and competitive analysis data have you gathered, and how does it influence your strategy?',

                # Qs for Market potential report
                # "What is the full legal name of your company, and what is its primary mission?",
                # "Can you describe the key product or technology your company has developed?",
                # "Who is the target audience for your product or service?",
                "What specific problem does your product or service solve for your target audience?",
                "How does your product or service stand out from existing market offerings?",
                "What pricing strategy has your company adopted for its product or service?",
                # "Could you explain your company's business model and how it generates revenue?",
                "What are the primary and potential secondary revenue streams for your company?",
                "How is your company's cost structure organized, and what impact does it have on pricing and profitability?",
                "Which sales and distribution channels is your company planning to use?",
                "Who are your main competitors, and what differentiates your product or service from theirs?"
                ]
# Specify the file path where you want to save the document
file_path = 'Questionnaire.docx'

# Create the questionnaire
create_questionnaire(questions, file_path)

print(f'Questionnaire has been saved to {file_path}')
